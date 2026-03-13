from langchain_community.vectorstores import Chroma
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument
import pandas as pd
from bs4 import BeautifulSoup
import xml.etree.ElementTree as ET
from pdf2image import convert_from_path
from mistralai import Mistral
from PIL import Image
import io
import base64
from typing import List, Dict
from dotenv import load_dotenv
import os
import re

# Load environment variables
load_dotenv()

# OpenRouter configuration
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1"
MISTRAL_API_KEY = os.getenv("MISTRAL_API_KEY")

# Embeddings using OpenRouter
embeddings = OpenAIEmbeddings(
    model="openai/text-embedding-3-large",
    openai_api_key=OPENROUTER_API_KEY,
    openai_api_base=OPENROUTER_BASE_URL
)

# Vector store setup
vectorstore = Chroma(
    persist_directory="./chroma_db",
    embedding_function=embeddings
)

# Add this line:
retriever = vectorstore.as_retriever(search_kwargs={"k": 5})

# Add this function at module level:
def format_docs(docs):
    """Format documents for context"""
    return "\n\n".join(doc.page_content for doc in docs)

# LLM with OpenRouter (using Claude by default)
llm = ChatOpenAI(
    model="anthropic/claude-3-5-sonnet",
    openai_api_key=OPENROUTER_API_KEY,
    openai_api_base=OPENROUTER_BASE_URL,
    temperature=0
)


def extract_text_from_file(file_path: str) -> str:
    """Extract text from various file formats"""
    ext = file_path.lower().split('.')[-1]
    
    try:
        if ext == 'txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        
        elif ext == 'pdf':
            # Try regular PDF extraction first
            try:
                loader = PyPDFLoader(file_path)
                docs = loader.load()
                text = '\n'.join([doc.page_content for doc in docs])
                
                # If extracted text is too short, likely scanned PDF
                if len(text.strip()) < 100:
                    print(f"  ⚠️  PDF appears to be scanned, using Mistral OCR...")
                    return extract_pdf_with_mistral_ocr(file_path)
                
                return text
            except Exception as e:
                print(f"  ⚠️  Standard PDF extraction failed, trying OCR: {str(e)}")
                return extract_pdf_with_mistral_ocr(file_path)
        
        elif ext == 'docx':
            doc = DocxDocument(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])
        
        elif ext == 'xlsx':
            df = pd.read_excel(file_path)
            return df.to_string()
        
        elif ext == 'csv':
            df = pd.read_csv(file_path)
            return df.to_string()
        
        elif ext == 'html':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                soup = BeautifulSoup(f, 'html.parser')
                return soup.get_text()
        
        elif ext == 'xml':
            tree = ET.parse(file_path)
            root = tree.getroot()
            # Extract all text from XML elements
            return ' '.join(root.itertext())
        
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    except Exception as e:
        raise Exception(f"Error extracting text from {file_path}: {str(e)}")


def extract_pdf_with_mistral_ocr(file_path: str) -> str:
    """Use Mistral Pixtral for OCR on image-based PDFs"""

    if not MISTRAL_API_KEY:
        raise Exception("MISTRAL_API_KEY not found in environment variables")

    # Convert PDF pages to images
    images = convert_from_path(file_path, dpi=200)  # Lower DPI for speed

    # Initialize Mistral client
    client = Mistral(api_key=MISTRAL_API_KEY)

    extracted_text = []

    for i, image in enumerate(images):
        print(f"    Processing page {i+1}/{len(images)} with Mistral OCR...")

        # Convert PIL Image to base64
        buffered = io.BytesIO()
        image.save(buffered, format="PNG")
        img_str = base64.b64encode(buffered.getvalue()).decode()

        # Use Mistral Pixtral for OCR
        response = client.chat.complete(
            model="pixtral-12b-2409",
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": "Extract all text from this document page. Return ONLY the text content, no descriptions or commentary."},
                    {"type": "image_url", "image_url": f"data:image/png;base64,{img_str}"}
                ]
            }]
        )

        extracted_text.append(response.choices[0].message.content)

    return "\n\n".join(extracted_text)


async def add_document(file_path: str, filename: str):
    """Add a document to the vector store"""
    try:
        # Extract text using multi-format handler
        print(f"[DEBUG] Starting extraction for: {filename}")
        text = extract_text_from_file(file_path)
        print(f"[DEBUG] Extracted text length: {len(text) if text else 0}")
        print(f"[DEBUG] First 100 chars: {text[:100] if text else 'EMPTY'}")

        if not text or len(text.strip()) < 10:
            print(f"[DEBUG] Text too short or empty, returning 0 chunks")
            return {
                "success": True,
                "message": "File uploaded but no text content extracted",
                "chunks_processed": 0
            }

        # Create Document object
        print(f"[DEBUG] Creating Document object...")
        doc = Document(page_content=text, metadata={"filename": filename})
        print(f"[DEBUG] Document created with page_content length: {len(doc.page_content)}")

        # Split into chunks
        print(f"[DEBUG] Splitting into chunks...")
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        chunks = text_splitter.split_documents([doc])
        print(f"[DEBUG] Created {len(chunks)} chunks")

        # Add to vector store
        print(f"[DEBUG] Adding {len(chunks)} chunks to vectorstore...")
        # Add documents in batches to avoid memory issues
        # Use smaller batch size (25) for large documents to prevent hangs
        batch_size = 25
        total_batches = (len(chunks) + batch_size - 1) // batch_size

        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            vectorstore.add_documents(batch)

        print(f"[DEBUG] Successfully added all {len(chunks)} chunks to vectorstore")

        return {
            "success": True,
            "message": f"Successfully added {filename}",
            "chunks_processed": len(chunks)
        }

    except Exception as e:
        print(f"[DEBUG] Exception occurred: {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        raise Exception(f"Error processing {filename}: {str(e)}")


async def add_document_from_supabase(document_text: str, metadata: Dict):
    """Add a document to the vector store from Supabase data
    
    Args:
        document_text: The full text content from Supabase 'document' field
        metadata: Dictionary containing all metadata fields:
            - id (uuid)
            - title (text)
            - summary (text)
            - key_terms (jsonb/text array)
            - main_topics (jsonb/text array)
            - geographic_scope (jsonb/text array)
            - effective_date (text/date)
            - created_at (timestamp)
            - Any other fields from the Supabase row
    
    Returns:
        dict with success status, message, and chunks_processed count
    """
    try:
        doc_id = metadata.get('id', 'unknown')
        title = metadata.get('title', 'Untitled Document')
        
        print(f"[INGEST] Processing document from Supabase: id={doc_id}, title={title[:50]}...")
        
        if not document_text or len(document_text.strip()) < 10:
            print(f"[INGEST] Document text too short or empty, returning 0 chunks")
            return {
                "success": True,
                "message": "Document text too short, skipping",
                "chunks_processed": 0
            }
        
        # Build comprehensive metadata for Chroma
        # Chroma metadata should be flat, so convert arrays/lists to JSON strings if needed
        chroma_metadata = {
            "document_id": str(doc_id),
            "title": title,
            "source": "supabase"
        }
        
        # Add optional metadata fields (convert lists/arrays to strings for Chroma compatibility)
        if metadata.get('summary'):
            chroma_metadata['summary'] = str(metadata['summary'])
        
        if metadata.get('key_terms'):
            key_terms = metadata['key_terms']
            if isinstance(key_terms, list):
                chroma_metadata['key_terms'] = ', '.join(str(term) for term in key_terms)
            else:
                chroma_metadata['key_terms'] = str(key_terms)
        
        if metadata.get('main_topics'):
            main_topics = metadata['main_topics']
            if isinstance(main_topics, list):
                chroma_metadata['main_topics'] = ', '.join(str(topic) for topic in main_topics)
            else:
                chroma_metadata['main_topics'] = str(main_topics)
        
        if metadata.get('geographic_scope'):
            geo_scope = metadata['geographic_scope']
            if isinstance(geo_scope, list):
                chroma_metadata['geographic_scope'] = ', '.join(str(scope) for scope in geo_scope)
            else:
                chroma_metadata['geographic_scope'] = str(geo_scope)
        
        if metadata.get('effective_date'):
            chroma_metadata['effective_date'] = str(metadata['effective_date'])
        
        if metadata.get('created_at'):
            chroma_metadata['created_at'] = str(metadata['created_at'])
        
        # Use title as filename for consistency with existing add_document function
        chroma_metadata['filename'] = title
        
        # Create Document object with full metadata
        doc = Document(page_content=document_text, metadata=chroma_metadata)
        print(f"[INGEST] Created document with {len(document_text)} chars")
        
        # Split into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        chunks = text_splitter.split_documents([doc])
        print(f"[INGEST] Created {len(chunks)} chunks")
        
        # Add to vector store in batches
        batch_size = 25
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            vectorstore.add_documents(batch)
        
        print(f"[INGEST] Successfully added {len(chunks)} chunks to vectorstore")
        
        return {
            "success": True,
            "message": f"Successfully added document: {title}",
            "chunks_processed": len(chunks)
        }
    
    except Exception as e:
        print(f"[INGEST] Exception occurred: {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        raise Exception(f"Error processing document from Supabase: {str(e)}")


def strip_markdown(text: str) -> str:
    """Strip markdown formatting from text to prevent formatting corruption in conversation history"""
    if not text:
        return text
    
    # Remove bold/italic markers (**text**, *text*, __text__, _text_)
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # **bold**
    text = re.sub(r'\*([^*]+)\*', r'\1', text)      # *italic* (but not **)
    text = re.sub(r'__([^_]+)__', r'\1', text)       # __bold__
    text = re.sub(r'_([^_]+)_', r'\1', text)        # _italic_ (but not __)
    
    # Remove headers (# ## ###)
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    
    # Remove links [text](url) -> text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    
    # Remove inline code `code`
    text = re.sub(r'`([^`]+)`', r'\1', text)
    
    # Remove code blocks ```code```
    text = re.sub(r'```[\s\S]*?```', '', text)
    
    # Clean up extra whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)  # Max 2 newlines
    text = text.strip()
    
    return text

def extract_key_terms_from_conversation(conversation_history: List[Dict]) -> str:
    """Extract key terms and entities from conversation history for query expansion"""
    if not conversation_history:
        return ""
    
    key_terms = []
    for entry in conversation_history:
        query = entry.get('user_query', '')
        answer = entry.get('final_answer', '')
        
        # Extract important terms (capitalized words, numbers, specific terms)
        import re
        # Find capitalized phrases (likely entities/terms)
        capitalized = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', query + ' ' + answer)
        # Find numbers (likely dates, costs, etc.)
        numbers = re.findall(r'\b\d+(?:\.\d+)?\b', query + ' ' + answer)
        # Find quoted phrases
        quoted = re.findall(r'"([^"]+)"', query + ' ' + answer)
        
        key_terms.extend(capitalized[:3])  # Limit to avoid too much expansion
        key_terms.extend(numbers[:2])
        key_terms.extend(quoted)
    
    # Remove duplicates and common words
    common_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'}
    unique_terms = [term for term in set(key_terms) if term.lower() not in common_words and len(term) > 2]
    
    return ' '.join(unique_terms[:10])  # Limit to 10 terms max

async def query_rag(question: str, conversation_history: List[Dict] = None):
    """Query the RAG system and return answer with source documents
    
    Args:
        question: The current question to answer
        conversation_history: Optional list of previous Q&A pairs from the session
    """
    # Build conversation context string
    conversation_context = ""
    if conversation_history:
        conv_lines = []
        for entry in conversation_history:
            user_q = entry.get('user_query', '')
            assistant_a = entry.get('final_answer', '')
            if user_q:
                conv_lines.append(f"User: {user_q}")
            if assistant_a:
                # Strip markdown to prevent formatting corruption, then truncate
                clean_answer = strip_markdown(assistant_a)
                truncated_answer = clean_answer[:300] + "..." if len(clean_answer) > 300 else clean_answer
                conv_lines.append(f"Assistant: {truncated_answer}")
        
        if conv_lines:
            conversation_context = "\n\nPrevious conversation:\n" + "\n".join(conv_lines)
    
    # Expand query for better retrieval using key terms from conversation
    expanded_query = question
    if conversation_history:
        key_terms = extract_key_terms_from_conversation(conversation_history)
        if key_terms:
            expanded_query = f"{question} {key_terms}"
            print(f"   📝 Expanded query with conversation context: {key_terms[:50]}...")
    
    # Build prompt template with optional conversation context
    if conversation_context:
        template = """Answer the question based on the following context, considering the previous conversation:

{conversation_context}

Context from documents: {context}

Current question: {question}

Answer the question directly, considering the conversation context. Cite the specific document sources used."""
    else:
        template = """Answer the question based on the following context:

Context: {context}

Question: {question}

Answer the question directly and cite the specific document sources used."""

    prompt = ChatPromptTemplate.from_template(template)

    # Retrieve documents using expanded query
    # Note: retriever.invoke() is synchronous, but we're in an async function
    # This is fine as long as the operation is fast
    docs = retriever.invoke(expanded_query)
    
    # Extract source filenames from document metadata BEFORE formatting
    sources = []
    for doc in docs:
        if hasattr(doc, 'metadata') and doc.metadata:
            filename = doc.metadata.get('filename', 'Unknown')
            if filename and filename not in sources:
                sources.append(filename)
    
    # Format context from documents
    context = format_docs(docs)
    
    # Create chain with pre-formatted context
    # Use RunnableLambda to inject the context
    from langchain_core.runnables import RunnableLambda
    
    def create_context_dict(q):
        """Create dict with context, question, and conversation history"""
        if conversation_context:
            return {
                "conversation_context": conversation_context,
                "context": context,
                "question": q
            }
        else:
            return {"context": context, "question": q}
    
    chain = (
        RunnableLambda(create_context_dict)
        | prompt
        | llm
        | StrOutputParser()
    )

    response = chain.invoke(question)

    return {
        "answer": response,
        "sources": sources,
        "source_documents": docs  # Include full documents for reference
    }


def list_documents():
    """List all documents in the vector store with structured metadata
    
    Returns documents grouped by source with metadata for easy frontend rendering.
    """
    from collections import defaultdict
    
    # Get all documents
    docs = vectorstore.get()
    
    if not docs or not docs.get('ids'):
        return {
            "total_chunks": 0,
            "unique_documents": 0,
            "sources": {},
            "documents": []  # Legacy format for backwards compatibility
        }
    
    metadatas = docs.get('metadatas', [])
    ids = docs.get('ids', [])
    total_chunks = len(ids)
    
    # Group documents by source and aggregate metadata
    documents_by_source = defaultdict(lambda: {
        'documents': {},
        'total_chunks': 0,
        'unique_document_count': 0
    })
    
    # Process all chunks to aggregate by document
    for i, metadata in enumerate(metadatas):
        source = metadata.get('source', 'file_upload')  # Default for legacy documents
        title = metadata.get('title') or metadata.get('filename', 'Untitled Document')
        doc_id = metadata.get('document_id', title)  # Use title as ID if no document_id
        
        # Create a unique key for this document
        doc_key = (title, doc_id)
        
        # Initialize document entry if first time seeing it
        if doc_key not in documents_by_source[source]['documents']:
            documents_by_source[source]['documents'][doc_key] = {
                'title': title,
                'document_id': doc_id,
                'chunk_count': 0,
                'metadata': {
                    'summary': metadata.get('summary'),
                    'key_terms': metadata.get('key_terms'),
                    'main_topics': metadata.get('main_topics'),
                    'geographic_scope': metadata.get('geographic_scope'),
                    'effective_date': metadata.get('effective_date'),
                    'created_at': metadata.get('created_at')
                }
            }
            documents_by_source[source]['unique_document_count'] += 1
        
        # Increment chunk count
        documents_by_source[source]['documents'][doc_key]['chunk_count'] += 1
        documents_by_source[source]['total_chunks'] += 1
    
    # Convert to list format for JSON serialization
    sources_dict = {}
    all_documents = []  # Legacy flat list
    
    for source, source_data in documents_by_source.items():
        # Convert documents dict to sorted list
        doc_list = []
        for doc_key, doc_info in sorted(source_data['documents'].items(), key=lambda x: x[1]['title']):
            doc_list.append({
                'title': doc_info['title'],
                'document_id': doc_info['document_id'],
                'chunk_count': doc_info['chunk_count'],
                'metadata': doc_info['metadata']
            })
            all_documents.append(doc_info['title'])  # Add to legacy list
        
        sources_dict[source] = {
            'documents': doc_list,
            'total_chunks': source_data['total_chunks'],
            'unique_document_count': source_data['unique_document_count']
        }
    
    return {
        "total_chunks": total_chunks,
        "unique_documents": sum(s['unique_document_count'] for s in sources_dict.values()),
        "sources": sources_dict,
        "documents": sorted(set(all_documents))  # Legacy format for backwards compatibility
    }
